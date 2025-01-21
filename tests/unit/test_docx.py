import unittest
import os
from pathlib import Path
import logging
import docx

class TestDocxProcessing(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        # Configure logging
        logging.basicConfig(level=logging.INFO)
        cls.logger = logging.getLogger('DocxTester')
        
        # Get the project root directory
        cls.project_root = Path(os.path.dirname(os.path.abspath(__file__))).parent.parent
        cls.test_files_dir = os.path.join(cls.project_root, "test_files")
        
        cls.logger.info(f"Project root: {cls.project_root}")
        cls.logger.info(f"Test files directory: {cls.test_files_dir}")

    def test_directory_exists(self):
        """Test if test_files directory exists"""
        self.assertTrue(os.path.exists(self.test_files_dir), 
                       f"Directory does not exist: {self.test_files_dir}")
        
        # List directory contents
        contents = os.listdir(self.test_files_dir)
        self.logger.info("Files in test_files directory:")
        for file in contents:
            self.logger.info(f"- {file}")

    def test_find_docx_files(self):
        """Test finding DOCX files in directory"""
        docx_files = [f for f in os.listdir(self.test_files_dir) if f.endswith('.docx')]
        self.logger.info(f"Found DOCX files: {docx_files}")
        self.assertGreater(len(docx_files), 0, "No DOCX files found in directory")

    def test_read_docx_file(self):
        """Test reading DOCX file content"""
        docx_files = [f for f in os.listdir(self.test_files_dir) if f.endswith('.docx')]
        for docx_file in docx_files:
            docx_path = os.path.join(self.test_files_dir, docx_file)
            self.logger.info(f"Testing file: {docx_path}")
            
            # Check file exists
            self.assertTrue(os.path.exists(docx_path), 
                          f"File does not exist: {docx_path}")
            
            # Check file size
            size = os.path.getsize(docx_path)
            self.assertGreater(size, 0, 
                             f"File is empty: {docx_path}")
            
            # Try to read file
            try:
                doc = docx.Document(docx_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                self.logger.info(f"Successfully read: {docx_file}")
                self.logger.info(f"Number of paragraphs: {len(doc.paragraphs)}")
                self.logger.info(f"First 100 chars: {text[:100]}")
                
                # Basic content checks
                self.assertGreater(len(doc.paragraphs), 0, 
                                 f"No paragraphs found in {docx_file}")
                self.assertGreater(len(text), 0, 
                                 f"No text content found in {docx_file}")
                
            except Exception as e:
                self.fail(f"Failed to read {docx_file}: {str(e)}")

if __name__ == '__main__':
    unittest.main(verbosity=2)
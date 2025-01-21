import PyPDF2
import docx
import os
from typing import Dict, Optional, Union, Tuple  # Changed import
import logging
from datetime import datetime

class DocumentProcessor:
    """
    Handles processing of CV documents in various formats (PDF, DOCX).
    """
    
    def __init__(self):
        self._setup_logger()
        self.supported_formats = {'.pdf', '.docx', '.doc'}

    def _setup_logger(self):
        self.logger = logging.getLogger('DocumentProcessor')
        self.logger.setLevel(logging.INFO)
        
        os.makedirs('logs', exist_ok=True)
        
        file_handler = logging.FileHandler(
            f'logs/document_processor_{datetime.now().strftime("%Y%m%d")}.log'
        )
        console_handler = logging.StreamHandler()
        
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)
        
        self.logger.addHandler(file_handler)
        self.logger.addHandler(console_handler)

    def process_document(self, file_path: str) -> Dict[str, Union[str, bool, Dict]]:
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            if file_ext == '.pdf':
                content, metadata = self._process_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                content, metadata = self._process_docx(file_path)
            else:
                raise ValueError(f"Unexpected file format: {file_ext}")
            
            self.logger.info(f"Successfully processed document: {file_path}")
            
            return {
                'success': True,
                'content': content,
                'metadata': metadata,
                'error': None
            }
            
        except Exception as e:
            error_msg = f"Error processing document: {str(e)}"
            self.logger.error(error_msg)
            return {
                'success': False,
                'content': None,
                'metadata': None,
                'error': error_msg
            }

    def _process_pdf(self, file_path: str) -> Tuple[str, Dict]:  # Changed return type hint
        self.logger.info(f"Processing PDF: {file_path}")
        
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            content = ""
            for page in reader.pages:
                content += page.extract_text() + "\n"
            
            metadata = {
                'pages': len(reader.pages),
                'encrypted': reader.is_encrypted,
                'file_size': os.path.getsize(file_path)
            }
            
            return content.strip(), metadata

    def _process_docx(self, file_path: str) -> Tuple[str, Dict]:  # Changed return type hint
        self.logger.info(f"Processing DOCX: {file_path}")
        
        doc = docx.Document(file_path)
        
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        metadata = {
            'paragraphs': len(doc.paragraphs),
            'sections': len(doc.sections),
            'file_size': os.path.getsize(file_path)
        }
        
        return content.strip(), metadata

    def validate_document(self, file_path: str) -> Dict[str, Union[bool, str]]:
        try:
            if not os.path.exists(file_path):
                return {
                    'valid': False,
                    'error': 'File does not exist'
                }
            
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext not in self.supported_formats:
                return {
                    'valid': False,
                    'error': f'Unsupported file format: {file_ext}'
                }
            
            if os.path.getsize(file_path) > 10 * 1024 * 1024:
                return {
                    'valid': False,
                    'error': 'File size exceeds 10MB limit'
                }
            
            return {
                'valid': True,
                'error': None
            }
            
        except Exception as e:
            return {
                'valid': False,
                'error': f'Validation error: {str(e)}'
            }